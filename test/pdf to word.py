# -*- coding: utf-8 -*-
from pdfminer.pdfparser import PDFParser, PDFDocument
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.layout import LAParams
from pdfminer.converter import PDFPageAggregator
from pdfminer.pdfinterp import PDFTextExtractionNotAllowed
from docx import Document
import os

document = Document()


def createWord(wordpath, pdfpath):
    # rb以二进制读模式打开本地pdf文件
    fn = open(pdfpath, 'rb')
    # 创建一个pdf文档分析器
    parser = PDFParser(fn)
    # 创建一个PDF文档
    doc = PDFDocument()
    # 连接分析器 与文档对象
    parser.set_document(doc)
    doc.set_parser(parser)

    # 提供初始化密码doc.initialize("lianxipython")
    # 如果没有密码 就创建一个空的字符串
    doc.initialize("")
    # 检测文档是否提供txt转换，不提供就忽略
    if not doc.is_extractable:
        raise PDFTextExtractionNotAllowed

    else:
        # 创建PDf资源管理器
        resource = PDFResourceManager()
        # 创建一个PDF参数分析器
        laparams = LAParams()
        # 创建聚合器,用于读取文档的对象
        device = PDFPageAggregator(resource, laparams=laparams)
        # 创建解释器，对文档编码，解释成Python能够识别的格式
        interpreter = PDFPageInterpreter(resource, device)
        # 循环遍历列表，每次处理一页的内容
        # doc.get_pages() 获取page列表
        for page in doc.get_pages():
            # 利用解释器的process_page()方法解析读取单独页数
            interpreter.process_page(page)
            # 使用聚合器get_result()方法获取内容
            layout = device.get_result()
            # 这里layout是一个LTPage对象,里面存放着这个page解析出的各种对象
            for out in layout:
                # 判断是否含有get_text()方法，获取我们想要的文字
                if hasattr(out, "get_text"):
                    # print(out.get_text(), type(out.get_text()))
                    content = out.get_text().replace(u'\xa0', u' ')  # 将'\xa0'替换成u' '空格，这个\xa0就是&nbps空格
                    # with open('test.txt','a') as f:
                    #     f.write(out.get_text().replace(u'\xa0', u' ')+'\n')
                    document.add_paragraph(
                        content, style='ListBullet'  # 添加段落，样式为unordered list类型
                    )
                document.save(wordpath)  # 保存这个文档


# 遍历当前目录，并把Pdf文件转换为Word
def pdfToWord():
    print("转换中...")
    # 获取当前运行路径
    path = os.getcwd()
    # 获取所有文件名的列表
    filename_list = os.listdir(path)
    # 获取所有pdf文件名列表
    pdfname_list = [filename for filename in filename_list \
                    if filename.endswith((".pdf"))]
    for pdfname in pdfname_list:
        # 分离pdf文件名称和后缀，转化为word名称
        wordname = os.path.splitext(pdfname)[0] + '.docx'
        # 如果当前pdf文件对应的word文件存在，则不转化
        if wordname in filename_list:
            continue
        # 拼接 路径和文件名
        wordpath = os.path.join(path, wordname)
        pdfpath = os.path.join(path, pdfname)
        createWord(wordpath, pdfpath)


if __name__ == '__main__':
    pdfToWord()